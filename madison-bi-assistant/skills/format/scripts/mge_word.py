"""
MGE Word Document Builder -- branded Word generation for Madison Group Enterprises.

Provides helper functions that create fully branded, print-ready Word documents
using python-docx. All styling boilerplate and OOXML gotchas are handled here
so the calling code focuses purely on content.

Design rationale:
- Tables over text boxes for ALL layout (KPI panels, callouts, metadata).
  Text boxes require massive OOXML injection and render inconsistently.
- Custom paragraph styles (not built-in Heading 1-9) to avoid theme font
  references that override font.name at the oxml layer.
- New OxmlElement per cell for shading -- an element can only have one parent.
- Field codes via w:fldChar complex fields (not w:fldSimple) for reliability.
- Section breaks with explicit dimension swap for landscape orientation.

Usage:
    import mge_word as word

    doc = word.create_branded_document("Quarterly Review", subtitle="Q2 FY2026")
    word.add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph("Body text here.", style="MGE Body")
    word.add_data_table(doc, ["Metric", "Value"], [["Revenue", "$12.4M"]])
    doc.save("output.docx")

Requires: python-docx >= 1.1.0, mge_brand.py (same directory)
"""

from __future__ import annotations

import os
import sys
from datetime import date
from typing import Optional, Sequence

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, Emu, RGBColor, Mm, Twips
from docx.table import Table, _Cell, _Row

# Ensure mge_brand is importable from the same directory.
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
if _SCRIPT_DIR not in sys.path:
    sys.path.insert(0, _SCRIPT_DIR)

import mge_brand as brand

__version__ = "1.0.0"


# ---------------------------------------------------------------------------
# Internal helpers -- OOXML primitives
# ---------------------------------------------------------------------------

def _rgb(colour: brand.Colour) -> RGBColor:
    """Convert a brand Colour to python-docx RGBColor."""
    return RGBColor(*colour.rgb)


def _clear_theme_fonts(run_element):
    """Remove theme font references from a run so font.name sticks in Word.

    Built-in heading styles carry w:rFonts theme attributes that silently
    override font.name. Our custom styles avoid built-in bases, but this
    is a safety net for any run that might inherit theme references.
    """
    rPr = run_element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
                rFonts.attrib.pop(qn(attr), None)


def _apply_type_level(font_obj, level: brand.TypeLevel):
    """Apply a TypeLevel's properties to a python-docx Font object."""
    font_obj.name = brand.FONT_NAME
    font_obj.size = Pt(level.size_pt)
    font_obj.bold = level.bold
    font_obj.color.rgb = _rgb(level.colour)


def _apply_para_spacing(para_format, level: brand.TypeLevel):
    """Apply spacing from a TypeLevel to a ParagraphFormat."""
    if level.space_before_pt:
        para_format.space_before = Pt(level.space_before_pt)
    else:
        para_format.space_before = Pt(0)
    para_format.space_after = Pt(level.space_after_pt)
    para_format.line_spacing = level.line_spacing


# ---------------------------------------------------------------------------
# Cell-level OOXML utilities
# ---------------------------------------------------------------------------

def set_cell_shading(cell: _Cell, colour_hex: str):
    """Set cell background fill. Creates a new element per cell (required).

    The colour_hex should be 6-char hex without '#' prefix, e.g. '3F5364'.
    Removes any existing w:shd to avoid duplicates.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), colour_hex)
    tcPr.append(shd)


def set_cell_borders(cell: _Cell, **kwargs):
    """Set per-edge borders on a cell.

    Keyword arguments are edge names (top, bottom, start, end) mapped to
    dicts with keys: val, sz, color, space.

    Example:
        set_cell_borders(cell,
            top={"val": "single", "sz": "8", "color": "3F5364", "space": "0"},
            bottom={"val": "single", "sz": "8", "color": "3F5364", "space": "0"},
        )

    sz is in eighths of a point: 4 = 0.5pt, 8 = 1pt, 12 = 1.5pt, 16 = 2pt.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcBorders'))
    if existing is not None:
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for edge, props in kwargs.items():
        el = OxmlElement(f'w:{edge}')
        for attr, val in props.items():
            el.set(qn(f'w:{attr}'), str(val))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def set_cell_padding(cell: _Cell, top: int = 72, bottom: int = 72,
                     start: int = 115, end: int = 115):
    """Set cell margins in dxa (twips). 72 dxa ~ 0.05in, 115 dxa ~ 0.08in.

    Removes any existing w:tcMar before appending to avoid duplicates.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('w:top', top), ('w:bottom', bottom),
                      ('w:start', start), ('w:end', end)]:
        el = OxmlElement(edge)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_table_borders(table: Table, **kwargs):
    """Set table-level borders. Same kwargs format as set_cell_borders."""
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for edge, props in kwargs.items():
        el = OxmlElement(f'w:{edge}')
        for attr, val in props.items():
            el.set(qn(f'w:{attr}'), str(val))
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_table_width_pct(table: Table, pct: int = 5000):
    """Set table width as percentage. 5000 = 100% of available space."""
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn('w:tblW'))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(pct))
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)


def _set_repeat_header(row: _Row):
    """Mark a table row to repeat as header on every page."""
    trPr = row._tr.get_or_add_trPr()
    existing = trPr.find(qn('w:tblHeader'))
    if existing is None:
        trPr.append(OxmlElement('w:tblHeader'))


def _remove_all_borders(table: Table):
    """Remove all borders from a table (for layout tables)."""
    none_border = {"val": "none", "sz": "0", "color": "auto", "space": "0"}
    _set_table_borders(table,
                       top=none_border, bottom=none_border,
                       start=none_border, end=none_border,
                       insideH=none_border, insideV=none_border)


def _set_paragraph_border(paragraph, edge: str, val: str = "single",
                          sz: str = "8", color: str = "3F5364",
                          space: str = "4"):
    """Add a border to a paragraph (e.g. bottom rule line for headers)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    el = OxmlElement(f'w:{edge}')
    el.set(qn('w:val'), val)
    el.set(qn('w:sz'), sz)
    el.set(qn('w:color'), color)
    el.set(qn('w:space'), space)
    pBdr.append(el)


# ---------------------------------------------------------------------------
# Field codes
# ---------------------------------------------------------------------------

def add_field_code(paragraph, field_type: str = 'PAGE'):
    """Insert a Word field code (PAGE, NUMPAGES, DATE, FILENAME) into a paragraph.

    Uses the w:fldChar complex field pattern (begin -> instrText -> separate -> end)
    which survives Word re-saving. Returns the run containing the field so the
    caller can style it.
    """
    run = paragraph.add_run()
    # Begin
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)
    # Instruction text -- leading/trailing spaces are mandatory for Word's parser
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' {field_type} '
    run._r.append(instrText)
    # Separate
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar_sep)
    # End
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)
    return run


# ---------------------------------------------------------------------------
# Style registration
# ---------------------------------------------------------------------------

def _register_paragraph_style(doc: DocumentType, name: str,
                              level: brand.TypeLevel, *,
                              alignment=None, outline_level: int | None = None):
    """Create and register a custom paragraph style.

    Uses WD_STYLE_TYPE.PARAGRAPH with no built-in base to avoid theme font
    inheritance. Sets font, spacing, and optionally outline level for TOC.
    """
    if name in [s.name for s in doc.styles]:
        return doc.styles[name]
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH, builtin=False)
    _apply_type_level(style.font, level)
    # Force Arial on the underlying XML to guarantee no theme font override
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), brand.FONT_NAME)
    rFonts.set(qn('w:hAnsi'), brand.FONT_NAME)
    rFonts.set(qn('w:cs'), brand.FONT_NAME)
    # Clear any theme attributes
    for attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
        rFonts.attrib.pop(qn(attr), None)

    _apply_para_spacing(style.paragraph_format, level)
    if alignment is not None:
        style.paragraph_format.alignment = alignment
    style.hidden = False
    style.quick_style = True
    # Outline level for navigation pane / TOC recognition
    if outline_level is not None:
        pPr = style.element.get_or_add_pPr()
        outline_el = OxmlElement('w:outlineLvl')
        outline_el.set(qn('w:val'), str(outline_level))
        pPr.append(outline_el)
    return style


def _register_character_style(doc: DocumentType, name: str,
                              colour: brand.Colour, bold: bool = False):
    """Create and register a custom character (run-level) style."""
    if name in [s.name for s in doc.styles]:
        return doc.styles[name]
    style = doc.styles.add_style(name, WD_STYLE_TYPE.CHARACTER, builtin=False)
    style.font.name = brand.FONT_NAME
    style.font.color.rgb = _rgb(colour)
    style.font.bold = bold
    style.hidden = False
    return style


def _register_all_styles(doc: DocumentType):
    """Register every MGE custom style on the document."""
    # Paragraph styles
    # Title: 30pt Bold Dark Grey, generous space_before for cover positioning
    title_level = brand.TypeLevel(30, True, brand.DARK_GREY, 0, 12, 1.0)
    _register_paragraph_style(doc, 'MGE Title', title_level)

    _register_paragraph_style(doc, 'MGE H1', brand.WORD_H1, outline_level=0)
    _register_paragraph_style(doc, 'MGE H2', brand.WORD_H2, outline_level=1)
    _register_paragraph_style(doc, 'MGE H3', brand.WORD_H3, outline_level=2)

    body_level = brand.TypeLevel(10, False, brand.CONNECT_GREY, 0, 6, 1.15)
    _register_paragraph_style(doc, 'MGE Body', body_level)

    body_bold_level = brand.TypeLevel(10, True, brand.CONNECT_GREY, 0, 6, 1.15)
    _register_paragraph_style(doc, 'MGE Body Bold', body_bold_level)

    caption_level = brand.TypeLevel(8, False, brand.CONNECT_GREY_50, 0, 3, 1.15)
    _register_paragraph_style(doc, 'MGE Caption', caption_level)

    footnote_level = brand.TypeLevel(7, False, brand.CONNECT_GREY_50, 0, 2, 1.15)
    _register_paragraph_style(doc, 'MGE Footnote', footnote_level)

    _register_paragraph_style(doc, 'MGE Subtitle', brand.WORD_SUBTITLE)

    meta_label = brand.TypeLevel(10, False, brand.CONNECT_GREY_50, 0, 2, 1.15)
    _register_paragraph_style(doc, 'MGE Metadata Label', meta_label)

    meta_value = brand.TypeLevel(11, False, brand.CONNECT_GREY, 0, 2, 1.15)
    _register_paragraph_style(doc, 'MGE Metadata Value', meta_value)

    # Character styles
    _register_character_style(doc, 'MGE Accent', brand.ACCENT_RED, bold=True)
    _register_character_style(doc, 'MGE Negative', brand.ACCENT_RED, bold=False)


# ---------------------------------------------------------------------------
# Section and page setup
# ---------------------------------------------------------------------------

def _configure_section(section, orientation: str = 'portrait'):
    """Set A4 dimensions, margins, and header/footer distances on a section."""
    if orientation == 'landscape':
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Twips(brand.WORD_PAGE_HEIGHT_DXA)
        section.page_height = Twips(brand.WORD_PAGE_WIDTH_DXA)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Twips(brand.WORD_PAGE_WIDTH_DXA)
        section.page_height = Twips(brand.WORD_PAGE_HEIGHT_DXA)
    section.top_margin = Inches(brand.WORD_MARGIN_INCHES)
    section.bottom_margin = Inches(brand.WORD_MARGIN_INCHES)
    section.left_margin = Inches(brand.WORD_MARGIN_INCHES)
    section.right_margin = Inches(brand.WORD_MARGIN_INCHES)
    section.header_distance = Twips(brand.WORD_HEADER_DIST_DXA)
    section.footer_distance = Twips(brand.WORD_FOOTER_DIST_DXA)


# ---------------------------------------------------------------------------
# Headers and footers
# ---------------------------------------------------------------------------

def add_branded_header(section, title: str):
    """Add branded header: document title left-aligned in CG 50% 9pt with rule line."""
    header = section.header
    header.is_linked_to_previous = False
    # Clear any existing content
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0]
    run = hp.add_run(title)
    run.font.name = brand.FONT_NAME
    run.font.size = Pt(9)
    run.font.color.rgb = _rgb(brand.CONNECT_GREY_50)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(4)
    # 1pt bottom border rule line in Connect Grey
    _set_paragraph_border(hp, 'bottom', val='single', sz='8',
                          color=brand.CONNECT_GREY.hex, space='4')


def add_branded_footer(section, title: str):
    """Add branded footer with title left, company centre, Page X of Y right.

    Uses tab stops for the three-zone layout. Field codes for page numbering.
    Top border rule line above the footer.
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()
    fp = footer.paragraphs[0]

    # Determine content width for tab stops
    page_width = section.page_width or Twips(brand.WORD_PAGE_WIDTH_DXA)
    left_margin = section.left_margin or Inches(brand.WORD_MARGIN_INCHES)
    right_margin = section.right_margin or Inches(brand.WORD_MARGIN_INCHES)
    content_width = page_width - left_margin - right_margin
    center_pos = content_width // 2

    # Add tab stops: centre and right
    tab_stops = fp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(center_pos, WD_ALIGN_PARAGRAPH.CENTER)
    tab_stops.add_tab_stop(content_width, WD_ALIGN_PARAGRAPH.RIGHT)

    _footer_font_size = Pt(8)
    _footer_colour = _rgb(brand.CONNECT_GREY_50)

    # Left zone: document title
    run_left = fp.add_run(title)
    run_left.font.name = brand.FONT_NAME
    run_left.font.size = _footer_font_size
    run_left.font.color.rgb = _footer_colour

    # Centre zone: company name
    run_tab1 = fp.add_run('\t')
    run_centre = fp.add_run('Madison Group Enterprises')
    run_centre.font.name = brand.FONT_NAME
    run_centre.font.size = _footer_font_size
    run_centre.font.color.rgb = _footer_colour

    # Right zone: Page X of Y
    run_tab2 = fp.add_run('\t')
    run_page_prefix = fp.add_run('Page ')
    run_page_prefix.font.name = brand.FONT_NAME
    run_page_prefix.font.size = _footer_font_size
    run_page_prefix.font.color.rgb = _footer_colour

    page_run = add_field_code(fp, 'PAGE')
    page_run.font.name = brand.FONT_NAME
    page_run.font.size = _footer_font_size
    page_run.font.color.rgb = _footer_colour

    run_of = fp.add_run(' of ')
    run_of.font.name = brand.FONT_NAME
    run_of.font.size = _footer_font_size
    run_of.font.color.rgb = _footer_colour

    numpages_run = add_field_code(fp, 'NUMPAGES')
    numpages_run.font.name = brand.FONT_NAME
    numpages_run.font.size = _footer_font_size
    numpages_run.font.color.rgb = _footer_colour

    # Top border rule line above footer
    _set_paragraph_border(fp, 'top', val='single', sz='8',
                          color=brand.CONNECT_GREY.hex, space='4')


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def add_cover_page(doc: DocumentType, title: str, subtitle: str | None = None,
                   metadata: dict | None = None, logo_path: str | None = None):
    """Build a branded cover page.

    - Logo top-left if logo_path provided (~50mm width)
    - Corner device top-right (Accent Red L-shape via coloured table cell)
    - Title at ~40% page height via paragraph spacing
    - Subtitle below title
    - Metadata table (Date, Author, Classification, Version)
    - The cover page section has no header/footer

    Args:
        doc: The Document to add the cover to.
        title: Document title text.
        subtitle: Optional subtitle text.
        metadata: Dict of label->value pairs. Defaults to date/author/classification.
        logo_path: Path to PNG logo file. Skipped gracefully if missing/None.
    """
    section = doc.sections[0]
    # Enable different first page so cover has no header/footer
    section.different_first_page_header_footer = True

    # -- Logo and corner device row (invisible 1x2 table) --
    logo_table = doc.add_table(rows=1, cols=2)
    _remove_all_borders(logo_table)
    _set_table_width_pct(logo_table)
    logo_table.autofit = False

    logo_cell = logo_table.cell(0, 0)
    device_cell = logo_table.cell(0, 1)

    # Logo (left cell)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_para.paragraph_format.space_after = Pt(0)
    if logo_path and os.path.isfile(logo_path):
        try:
            run = logo_para.add_run()
            run.add_picture(logo_path, width=Mm(50))
        except Exception:
            # Graceful degradation: skip logo on any error
            pass
    set_cell_padding(logo_cell, top=0, bottom=0, start=0, end=0)

    # Corner device (right cell) -- L-shaped brand element, top-right
    device_para = device_cell.paragraphs[0]
    device_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    device_para.paragraph_format.space_after = Pt(0)
    device_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    set_cell_padding(device_cell, top=0, bottom=0, start=0, end=0)
    # Insert corner device PNG (15mm x 15mm L-shape in Accent Red)
    device_img = brand.logo_path("corner-device-red")
    if device_img and os.path.isfile(str(device_img)):
        try:
            run = device_para.add_run()
            run.add_picture(str(device_img), width=Mm(15))
        except Exception:
            pass  # Graceful degradation: no corner device if image fails

    # -- Title positioning: push to ~40% of page height --
    # Usable height is roughly 9.69 inches with 1-inch margins.
    # 40% of that is ~3.88 inches. Subtract space used by logo row (~0.6in).
    p_spacer = doc.add_paragraph()
    p_spacer.paragraph_format.space_before = Inches(2.8)
    p_spacer.paragraph_format.space_after = Pt(0)

    # Title
    p_title = doc.add_paragraph(style='MGE Title')
    run = p_title.add_run(title)
    _apply_type_level(run.font, brand.TypeLevel(30, True, brand.DARK_GREY))
    _clear_theme_fonts(run._r)
    p_title.paragraph_format.space_after = Pt(6)

    # Subtitle
    if subtitle:
        p_sub = doc.add_paragraph(style='MGE Subtitle')
        run = p_sub.add_run(subtitle)
        _apply_type_level(run.font, brand.WORD_SUBTITLE)
        p_sub.paragraph_format.space_after = Pt(24)

    # Accent rule line
    p_rule = doc.add_paragraph()
    p_rule.paragraph_format.space_after = Pt(16)
    _set_paragraph_border(p_rule, 'top', val='single', sz='12',
                          color=brand.ACCENT_RED.hex, space='1')

    # Metadata table (invisible 2-column layout)
    if metadata is None:
        metadata = {
            'Date': date.today().strftime('%d %B %Y'),
            'Author': 'Madison Group Enterprises',
            'Classification': 'Internal',
            'Version': '1.0',
        }
    meta_table = doc.add_table(rows=len(metadata), cols=2)
    _remove_all_borders(meta_table)
    meta_table.autofit = False
    for i, (label, value) in enumerate(metadata.items()):
        label_cell = meta_table.cell(i, 0)
        value_cell = meta_table.cell(i, 1)
        label_cell.width = Mm(35)
        # Label
        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.name = brand.FONT_NAME
        lr.font.size = Pt(10)
        lr.font.color.rgb = _rgb(brand.CONNECT_GREY_50)
        lp.paragraph_format.space_after = Pt(2)
        set_cell_padding(label_cell, top=36, bottom=36, start=0, end=72)
        # Value
        vp = value_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.name = brand.FONT_NAME
        vr.font.size = Pt(11)
        vr.font.color.rgb = _rgb(brand.CONNECT_GREY)
        vp.paragraph_format.space_after = Pt(2)
        set_cell_padding(value_cell, top=36, bottom=36, start=0, end=0)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def create_branded_document(title: str, subtitle: str | None = None,
                            author: str = "Madison Group Enterprises",
                            include_cover: bool = True) -> DocumentType:
    """Create a fully branded Word document with all styles and sections configured.

    Args:
        title: Document title (used on cover page and headers).
        subtitle: Optional subtitle for the cover page.
        author: Author name for metadata and core properties.
        include_cover: Whether to build a cover page. Default True.

    Returns:
        A python-docx Document with all MGE styles registered, sections
        configured, cover page built (if requested), and branded
        headers/footers on body sections.
    """
    doc = Document()

    # Set core properties
    doc.core_properties.title = title
    doc.core_properties.author = author

    # Configure default style (Normal)
    normal = doc.styles['Normal']
    normal.font.name = brand.FONT_NAME
    normal.font.size = Pt(10)
    normal.font.color.rgb = _rgb(brand.CONNECT_GREY)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    # Register all custom styles
    _register_all_styles(doc)

    # Configure first section (A4 portrait)
    section = doc.sections[0]
    _configure_section(section, 'portrait')

    if include_cover:
        # Resolve logo
        logo = brand.logo_path("mge-wide")
        logo_str = str(logo) if logo else None

        add_cover_page(doc, title, subtitle=subtitle, metadata={
            'Date': date.today().strftime('%d %B %Y'),
            'Author': author,
            'Classification': 'Internal',
            'Version': '1.0',
        }, logo_path=logo_str)

        # Section break to body content
        body_section = doc.add_section(WD_SECTION.NEW_PAGE)
        _configure_section(body_section, 'portrait')

        # Branded header and footer on the body section
        add_branded_header(body_section, title)
        add_branded_footer(body_section, title)
    else:
        # No cover -- apply header/footer to the only section
        add_branded_header(section, title)
        add_branded_footer(section, title)

    return doc


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------

def add_heading(doc: DocumentType, text: str, level: int = 1):
    """Add a branded heading. H1 and H2 get a red full stop appended.

    Uses MGE H1/H2/H3 styles (not built-in Heading 1-9) to avoid theme
    font issues. The red full stop is a subtle brand accent -- skipped if
    the heading already ends with punctuation.

    Args:
        doc: The Document to add the heading to.
        text: Heading text.
        level: 1, 2, or 3.

    Returns:
        The paragraph object.
    """
    style_map = {1: 'MGE H1', 2: 'MGE H2', 3: 'MGE H3'}
    style_name = style_map.get(level, 'MGE H1')

    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    # Ensure font overrides are explicit on the run too
    type_levels = {1: brand.WORD_H1, 2: brand.WORD_H2, 3: brand.WORD_H3}
    tl = type_levels.get(level, brand.WORD_H1)
    _apply_type_level(run.font, tl)
    _clear_theme_fonts(run._r)

    # Red full stop for H1 and H2 (brand accent)
    if level in (1, 2) and text and text[-1] not in '.!?:;,':
        dot_run = p.add_run('.')
        dot_run.font.name = brand.FONT_NAME
        dot_run.font.size = Pt(tl.size_pt)
        dot_run.font.bold = True
        dot_run.font.color.rgb = _rgb(brand.ACCENT_RED)

    return p


# ---------------------------------------------------------------------------
# Data tables
# ---------------------------------------------------------------------------

def add_data_table(doc: DocumentType, headers: list[str], rows: list[list[str]],
                   col_widths: list[float] | None = None,
                   style: str = 'primary') -> Table:
    """Add a branded data table.

    Primary style: Connect Grey header fill, white bold text, alternating
    CG 10%/white rows, horizontal borders only.
    Secondary style: CG 20% header fill with CG text.

    Args:
        doc: The Document.
        headers: Column header strings.
        rows: List of row data (each row is a list of strings).
        col_widths: Optional column widths in inches. If None, uses equal widths.
        style: 'primary' or 'secondary'.

    Returns:
        The Table object.
    """
    n_cols = len(headers)
    n_rows = len(rows)
    table = doc.add_table(rows=1 + n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    # Full-width table
    _set_table_width_pct(table)

    # Table-level borders: horizontal only
    cg_hex = brand.CONNECT_GREY.hex
    cg20_hex = brand.CONNECT_GREY_20.hex
    _set_table_borders(table,
        top={"val": "single", "sz": "8", "color": cg_hex, "space": "0"},
        bottom={"val": "single", "sz": "8", "color": cg_hex, "space": "0"},
        insideH={"val": "single", "sz": "4", "color": cg20_hex, "space": "0"},
        insideV={"val": "none", "sz": "0", "color": "auto", "space": "0"},
        start={"val": "none", "sz": "0", "color": "auto", "space": "0"},
        end={"val": "none", "sz": "0", "color": "auto", "space": "0"},
    )

    # Determine header colours based on style
    if style == 'secondary':
        hdr_fill = brand.CONNECT_GREY_20.hex
        hdr_colour = brand.CONNECT_GREY
    else:
        hdr_fill = brand.CONNECT_GREY.hex
        hdr_colour = brand.WHITE

    # Header row
    header_row = table.rows[0]
    header_row.height = Pt(20)
    header_row.height_rule = 2  # EXACTLY
    _set_repeat_header(header_row)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = brand.FONT_NAME
        run.font.size = Pt(brand.WORD_TABLE_HEADER.size_pt)
        run.font.bold = brand.WORD_TABLE_HEADER.bold
        run.font.color.rgb = _rgb(hdr_colour)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_cell_shading(cell, hdr_fill)
        set_cell_padding(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_obj = table.rows[r_idx + 1]
        row_obj.height = Pt(18)
        row_obj.height_rule = 1  # AT_LEAST
        fill_hex = brand.CONNECT_GREY_10.hex if r_idx % 2 == 1 else brand.WHITE.hex
        for c_idx, val in enumerate(row_data):
            cell = table.cell(r_idx + 1, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = brand.FONT_NAME
            run.font.size = Pt(brand.WORD_TABLE_BODY.size_pt)
            run.font.bold = False
            run.font.color.rgb = _rgb(brand.WORD_TABLE_BODY.colour)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            set_cell_shading(cell, fill_hex)
            set_cell_padding(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths and len(col_widths) == n_cols:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    elif n_cols > 0:
        # Equal widths filling content area
        content_in = brand.WORD_CONTENT_WIDTH_DXA / 1440.0
        equal_w = content_in / n_cols
        for i in range(n_cols):
            for row in table.rows:
                row.cells[i].width = Inches(equal_w)

    return table


# ---------------------------------------------------------------------------
# Total rows (within tables)
# ---------------------------------------------------------------------------

def add_total_row(table: Table, label: str, values: list[str],
                  style: str = 'subtotal'):
    """Add a total row to an existing table.

    Args:
        table: The Table object (row is appended).
        label: Text for the first column.
        values: Values for remaining columns.
        style: 'subtotal' (bold + single top border in CG 50%) or
               'grand' (bold + double bottom border in CG + single top border).

    Returns:
        The new row object.
    """
    # Add a new row to the table
    row_cells = table.add_row().cells
    all_vals = [label] + list(values)

    for i, val in enumerate(all_vals):
        if i >= len(row_cells):
            break
        cell = row_cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.name = brand.FONT_NAME
        run.font.size = Pt(brand.WORD_TABLE_TOTAL.size_pt)
        run.font.bold = True
        run.font.color.rgb = _rgb(brand.WORD_TABLE_TOTAL.colour)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        set_cell_padding(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        if style == 'grand':
            set_cell_borders(cell,
                top={"val": "single", "sz": "8", "color": brand.CONNECT_GREY.hex, "space": "0"},
                bottom={"val": "double", "sz": "8", "color": brand.CONNECT_GREY.hex, "space": "0"},
            )
            set_cell_shading(cell, brand.WHITE.hex)
        else:
            # Subtotal
            set_cell_borders(cell,
                top={"val": "single", "sz": "4", "color": brand.CONNECT_GREY_50.hex, "space": "0"},
            )
            set_cell_shading(cell, brand.WHITE.hex)

    return table.rows[-1]


# ---------------------------------------------------------------------------
# KPI panels
# ---------------------------------------------------------------------------

def add_kpi_panel(doc: DocumentType, metrics: list[dict]):
    """Add a KPI panel as a side-by-side row of metric cards.

    Each metric dict should have:
        label (str): Metric name, e.g. "Revenue"
        value (str): Display value, e.g. "$12.4M"
        change (str, optional): Change text, e.g. "+$1.2M"
        change_pct (str, optional): Percentage change, e.g. "+10.7%"

    Uses an invisible table for layout. Big number 24pt Bold Dark Grey,
    label 10pt CG 50%, trend with directional colour.

    Returns:
        The Table object.
    """
    n = len(metrics)
    if n == 0:
        return None

    table = doc.add_table(rows=1, cols=n)
    _remove_all_borders(table)
    _set_table_width_pct(table)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Equal column widths
    content_in = brand.WORD_CONTENT_WIDTH_DXA / 1440.0
    col_w = content_in / n

    for i, metric in enumerate(metrics):
        cell = table.cell(0, i)
        cell.text = ''
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Inches(col_w)

        # CG 10% fill
        set_cell_shading(cell, brand.CONNECT_GREY_10.hex)
        set_cell_padding(cell, top=144, bottom=144, start=115, end=115)

        # 2pt bottom accent border in Connect Grey
        set_cell_borders(cell,
            bottom={"val": "single", "sz": "16", "color": brand.CONNECT_GREY.hex, "space": "0"},
            top={"val": "none", "sz": "0", "color": "auto", "space": "0"},
            start={"val": "none", "sz": "0", "color": "auto", "space": "0"},
            end={"val": "none", "sz": "0", "color": "auto", "space": "0"},
        )

        # Big number
        p_value = cell.paragraphs[0]
        p_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_value.paragraph_format.space_after = Pt(2)
        run = p_value.add_run(metric.get('value', ''))
        _apply_type_level(run.font, brand.WORD_KPI_VALUE)
        # Override colour to Dark Grey for maximum impact
        run.font.color.rgb = _rgb(brand.DARK_GREY)

        # Label
        p_label = cell.add_paragraph()
        p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_label.paragraph_format.space_before = Pt(0)
        p_label.paragraph_format.space_after = Pt(2)
        run = p_label.add_run(metric.get('label', ''))
        _apply_type_level(run.font, brand.WORD_KPI_LABEL)

        # Trend indicator
        change = metric.get('change', '')
        change_pct = metric.get('change_pct', '')
        if change or change_pct:
            p_trend = cell.add_paragraph()
            p_trend.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_trend.paragraph_format.space_before = Pt(0)
            p_trend.paragraph_format.space_after = Pt(0)

            # Determine direction and colour
            trend_text = change_pct or change
            is_negative = trend_text.startswith('-') or trend_text.startswith(brand.TREND_DOWN)
            if is_negative:
                arrow = brand.TREND_DOWN
                colour = brand.ACCENT_RED
            elif trend_text.startswith('+') or trend_text.startswith(brand.TREND_UP):
                arrow = brand.TREND_UP
                colour = brand.CONNECT_GREY
            else:
                arrow = brand.TREND_FLAT
                colour = brand.CONNECT_GREY_50

            # Arrow
            ar = p_trend.add_run(f'{arrow} ')
            ar.font.name = brand.FONT_NAME
            ar.font.size = Pt(9)
            ar.font.bold = True
            ar.font.color.rgb = _rgb(colour)

            # Value text
            display = f'{change} {change_pct}'.strip() if change and change_pct else trend_text
            vr = p_trend.add_run(display)
            vr.font.name = brand.FONT_NAME
            vr.font.size = Pt(9)
            vr.font.bold = False
            vr.font.color.rgb = _rgb(colour)

    return table


# ---------------------------------------------------------------------------
# Callout boxes
# ---------------------------------------------------------------------------

def add_callout(doc: DocumentType, title: str, body: str,
                style: str = 'note') -> Table:
    """Add a branded callout box using a single-cell table.

    Styles:
        note    -- Satin Grey fill, Connect Grey left border
        warning -- Accent Red 20% fill, Accent Red left border
        insight -- CG 10% fill, Connect Grey left border

    Args:
        doc: The Document.
        title: Callout title text (bold).
        body: Callout body text.
        style: 'note', 'warning', or 'insight'.

    Returns:
        The Table object.
    """
    style_config = {
        'note': {
            'fill': brand.SATIN_GREY.hex,
            'border_colour': brand.CONNECT_GREY.hex,
        },
        'warning': {
            'fill': brand.ACCENT_RED_20.hex,
            'border_colour': brand.ACCENT_RED.hex,
        },
        'insight': {
            'fill': brand.CONNECT_GREY_10.hex,
            'border_colour': brand.CONNECT_GREY.hex,
        },
    }
    cfg = style_config.get(style, style_config['note'])

    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_width_pct(table)

    cell = table.cell(0, 0)
    cell.text = ''

    set_cell_shading(cell, cfg['fill'])
    set_cell_padding(cell, top=144, bottom=144, start=144, end=144)

    # Left accent border + suppress other borders
    set_cell_borders(cell,
        start={"val": "single", "sz": "24", "color": cfg['border_colour'], "space": "0"},
        top={"val": "nil", "sz": "0", "color": "auto", "space": "0"},
        bottom={"val": "nil", "sz": "0", "color": "auto", "space": "0"},
        end={"val": "nil", "sz": "0", "color": "auto", "space": "0"},
    )

    # Remove table-level borders
    _remove_all_borders(table)

    # Title paragraph
    p_title = cell.paragraphs[0]
    run = p_title.add_run(title)
    _apply_type_level(run.font, brand.WORD_CALLOUT_TITLE)
    p_title.paragraph_format.space_after = Pt(4)

    # Body paragraph
    p_body = cell.add_paragraph()
    run = p_body.add_run(body)
    _apply_type_level(run.font, brand.WORD_CALLOUT_BODY)
    p_body.paragraph_format.space_after = Pt(0)

    return table


# ---------------------------------------------------------------------------
# Section and page management
# ---------------------------------------------------------------------------

def add_section_break(doc: DocumentType, orientation: str = 'portrait'):
    """Add a section break with the specified orientation.

    Explicitly sets dimensions (swaps width/height for landscape), margins,
    and unlinks headers/footers so they can be configured independently.

    Returns:
        The new Section object.
    """
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    _configure_section(new_section, orientation)
    # Unlink so the new section can have its own header/footer
    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False
    return new_section


def add_page_break(doc: DocumentType):
    """Add a simple page break within the same section."""
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Module self-test
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    # Quick smoke test: create a branded document and save it
    out_dir = os.path.join(_SCRIPT_DIR, '..', '..', '..', '..', 'scratch')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, '_test_mge_word.docx')

    doc = create_branded_document(
        "Quarterly Business Review",
        subtitle="Q2 FY2026 -- Division 100",
        author="Business Intelligence",
    )

    add_heading(doc, "Executive Summary", level=1)
    doc.add_paragraph(
        "Revenue declined 12% quarter-on-quarter, driven primarily by a contraction "
        "in the Garland portfolio. Working capital remains tied up in excess inventory.",
        style='MGE Body',
    )

    add_heading(doc, "Key Performance Indicators", level=2)
    add_kpi_panel(doc, [
        {"label": "Revenue", "value": "$12.4M", "change": "+$1.2M", "change_pct": "+10.7%"},
        {"label": "Gross Margin", "value": "28.3%", "change": "-1.2pp", "change_pct": ""},
        {"label": "Inventory Turns", "value": "2.64x", "change": "", "change_pct": "-9.2%"},
    ])

    add_heading(doc, "Financial Summary", level=2)
    add_data_table(doc,
        headers=["Metric", "Q1 FY26", "Q2 FY26", "Variance"],
        rows=[
            ["Revenue", "$13.8M", "$12.4M", "-$1.4M"],
            ["COGS", "$9.9M", "$8.9M", "-$1.0M"],
            ["Gross Profit", "$3.9M", "$3.5M", "-$0.4M"],
        ],
        col_widths=[2.0, 1.5, 1.5, 1.5],
    )

    add_heading(doc, "Observations", level=2)
    add_callout(doc, "Key Insight", "Garland CNY buffer has not drawn down despite "
                "record February sales. Monitor closely.", style='insight')
    add_callout(doc, "Warning", "AKG Retail sales have plateaued for five consecutive "
                "months. Purchasing freeze recommended.", style='warning')

    # Landscape section test
    ls = add_section_break(doc, orientation='landscape')
    add_branded_header(ls, "Quarterly Business Review")
    add_branded_footer(ls, "Quarterly Business Review")
    add_heading(doc, "Detailed Financials (Landscape)", level=1)
    doc.add_paragraph("This section demonstrates landscape orientation.", style='MGE Body')

    doc.save(out_path)
    print(f"Test document saved: {out_path}")
